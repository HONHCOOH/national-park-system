"""Insert diagrams into thesis document using clean approach:
1. Find body indices BEFORE modifications
2. Create paragraphs at end  
3. Move from bottom to top using pre-computed indices"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PAPER_PATH = r'C:\dev\杨念祖-毕业论文-国家公园大模型系统的智能决策支持研究.docx'
DIAGRAMS_DIR = r'C:\dev\national-park-system\diagrams'
OUTPUT_PATH = r'C:\dev\杨念祖-毕业论文-国家公园大模型系统的智能决策支持研究_含图表.docx'

diagrams = [
    ('05_技术架构图.png',        '图2-1 系统技术栈架构图',   '2.5'),
    ('06_火灾防控业务流程图.png',  '图3-1 火灾防控业务流程图', '3.1.3'),
    ('01_系统架构图.png',         '图4-1 系统整体架构图',    '4.1 '),
    ('02_功能模块图.png',         '图4-2 系统功能模块图',    '4.2 '),
    ('03_ER图.png',              '图4-3 数据库E-R图',      '4.3 '),
    ('04_大模型集成架构图.png',    '图5-1 大模型集成架构图',  '5.3.4'),
]

print("Step 1: Loading document...")
doc = Document(PAPER_PATH)
body = doc.element.body

# Find all target paragraph elements and their body indices
targets = []  # (body_index, fpath, caption)
for fname, caption, section in diagrams:
    fpath = os.path.join(DIAGRAMS_DIR, fname)
    if not os.path.exists(fpath):
        print(f"  [SKIP] Missing: {fname}")
        continue
    found = False
    for p in doc.paragraphs:
        if p.text.strip().startswith(section):
            body_idx = list(body).index(p._element)
            targets.append((body_idx, fpath, caption))
            print(f"  Section '{section}' -> body idx {body_idx}")
            found = True
            break
    if not found:
        print(f"  [WARN] Section '{section}' not found")

# Step 2: Create all image+text paragraphs at the END
print("\nStep 2: Creating paragraphs at end...")
pairs = []  # (body_idx, img_element, cap_element)
for body_idx, img_path, caption in targets:
    # Image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(0)
    img_para.paragraph_format.space_after = Pt(2)
    img_run = img_para.add_run()
    img_run.add_picture(img_path, width=Inches(5.2))

    # Caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after = Pt(6)
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(9)
    rPr = cap_run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

    pairs.append((body_idx, img_para._element, cap_para._element))
    print(f"  Created pair for {caption}")

# Step 3: Move elements from end to target positions (bottom to top)
print("\nStep 3: Moving elements to correct positions...")
pairs.sort(key=lambda x: x[0], reverse=True)  # Largest body_idx first

for body_idx, img_elem, cap_elem in pairs:
    # body.insert() automatically moves the element if already a child
    body.insert(body_idx + 1, cap_elem)
    body.insert(body_idx + 1, img_elem)
    print(f"  Moved to body idx {body_idx}")

# Step 4: Save
print(f"\nStep 4: Saving to {OUTPUT_PATH}...")
doc.save(OUTPUT_PATH)

# Verify
print("\nVerifying...")
verify = Document(OUTPUT_PATH)
shapes = verify.inline_shapes
print(f"  Paragraphs: {len(verify.paragraphs)}, Inline shapes: {len(shapes)}")
if len(shapes) >= 6:
    print("  SUCCESS!")
else:
    print(f"  WARNING: Expected 6 images, found {len(shapes)}")
